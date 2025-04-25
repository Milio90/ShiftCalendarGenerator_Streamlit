import streamlit as st
import docx
from icalendar import Calendar, Event
from datetime import datetime, timedelta, date
import re
import os
import io
import tempfile
import platform
import shutil
import base64

st.set_page_config(
    page_title="Employee Shift Calendar Generator",
    page_icon="📅",
    layout="wide"
)

def read_docx_tables(uploaded_file):
    """Read all tables content from an uploaded DOCX file."""
    try:
        # Save the uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        doc = docx.Document(tmp_file_path)
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        
        if not doc.tables:
            st.warning("No tables found in the document.")
            return []
        
        tables_data = []
        
        for table_index, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                # Skip empty rows
                if any(row_data):
                    rows.append(row_data)
            
            tables_data.append(rows)
            st.write(f"Table {table_index+1}: Found {len(rows)} rows with data")
        
        return tables_data
    except Exception as e:
        st.error(f"Error reading document: {e}")
        return []

def parse_first_table(rows, month, year):
    """Parse the first table format (Regular and On-Call shifts) with month rollover detection."""
    shifts = []
    current_month = month
    current_year = year
    last_day = 0  # Track the last day number we've seen
    
    for row in rows:
        if len(row) < 4:  # Ensure row has enough columns
            continue
        
        try:
            # Extract day, month, day_of_week, and employees
            day = row[0].strip()
            day_of_week = row[2].strip() if len(row) > 2 else ""
            employees_cell = row[3].strip() if len(row) > 3 else ""
            
            # Skip header rows or rows without day number
            if not day.isdigit():
                continue
            
            day = int(day)
            
            # Check if we've rolled over to the next month
            # If current day is significantly less than the last day, we've likely
            # moved to the next month
            if day < last_day and last_day > 20 and day < 10:
                # Move to next month
                current_month += 1
                if current_month > 12:
                    current_month = 1
                    current_year += 1
                st.info(f"Month rollover detected: now processing {current_month}/{current_year}")
            
            last_day = day
            
            # Parse employee names (may contain two employees, one with asterisk)
            employees = employees_cell.split('\n')
            employees = [e.strip() for e in employees if e.strip()]
            
            for employee in employees:
                is_on_call = "*" in employee
                employee_name = employee.replace("*", "").strip()
                
                # Create shift date using current_month and current_year
                shift_date = date(current_year, current_month, day)
                
                shift_type = "On-Call Shift" if is_on_call else "Regular Shift"
                
                shifts.append({
                    'employee': employee_name,
                    'date': shift_date,
                    'day_of_week': day_of_week,
                    'shift_type': shift_type
                })
        except Exception as e:
            st.error(f"Error parsing row in first table {row}: {e}")
            continue
    
    return shifts

def parse_second_table(rows, month, year):
    """Parse the second table format (Μεγάλη, Μικρή, ΤΕΠ shifts) with month rollover detection."""
    shifts = []
    current_month = month
    current_year = year
    last_day = 0  # Track the last day number we've seen
    
    for row in rows:
        if len(row) < 6:  # Ensure row has enough columns for second table format
            continue
        
        try:
            # Extract day, month, day_of_week, and employees from different shifts
            day = row[0].strip()
            day_of_week = row[2].strip() if len(row) > 2 else ""
            megali_shift = row[3].strip() if len(row) > 3 else ""
            mikri_shift = row[4].strip() if len(row) > 4 else ""
            tep_shift = row[5].strip() if len(row) > 5 else ""
            
            # Skip header rows or rows without day number
            if not day.isdigit():
                continue
            
            day = int(day)
            
            # Check if we've rolled over to the next month
            # If current day is significantly less than the last day, we've likely
            # moved to the next month
            if day < last_day and last_day > 20 and day < 10:
                # Move to next month
                current_month += 1
                if current_month > 12:
                    current_month = 1
                    current_year += 1
                st.info(f"Month rollover detected: now processing {current_month}/{current_year}")
            
            last_day = day
            
            # Use current_month and current_year for the shift date
            shift_date = date(current_year, current_month, day)
            
            # Process Μεγάλη shift (24h)
            if megali_shift:
                employee_name = megali_shift.replace(">", "").strip()
                if employee_name:
                    shifts.append({
                        'employee': employee_name,
                        'date': shift_date,
                        'day_of_week': day_of_week,
                        'shift_type': "Μεγάλη Shift (24h)"
                    })
            
            # Process Μικρή shift (24h)
            if mikri_shift:
                employee_name = mikri_shift.replace(">", "").strip()
                if employee_name:
                    shifts.append({
                        'employee': employee_name,
                        'date': shift_date,
                        'day_of_week': day_of_week,
                        'shift_type': "Μικρή Shift (24h)"
                    })
            
            # Process ΤΕΠ shift (12h)
            if tep_shift:
                employee_name = tep_shift.replace(">", "").strip()
                if employee_name:
                    shifts.append({
                        'employee': employee_name,
                        'date': shift_date,
                        'day_of_week': day_of_week,
                        'shift_type': "TEP Shift (12h)"
                    })
                
        except Exception as e:
            st.error(f"Error parsing row in second table {row}: {e}")
            continue
    
    return shifts

def parse_specialty_on_call_table(rows):
    """Parse the specialty on-call table format with date (DD-MM-YYYY or DD/MM/YYYY) in first column."""
    shifts = []
    
    for row in rows:
        if len(row) < 3:  # Ensure row has enough columns
            continue
        
        try:
            # Extract date, day_of_week, and employee
            date_str = row[0].strip()
            day_of_week = row[1].strip() if len(row) > 1 else ""
            employee_name = row[2].strip() if len(row) > 2 else ""
            
            # Skip header rows or rows without proper date format
            # Updated regex to match both DD-MM-YYYY and DD/MM/YYYY formats
            if not re.match(r"\d{1,2}[-/]\d{1,2}[-/]\d{4}", date_str):
                continue
            
            # Parse date (supports both DD-MM-YYYY and DD/MM/YYYY)
            if '-' in date_str:
                day, month, year = map(int, date_str.split('-'))
            elif '/' in date_str:
                day, month, year = map(int, date_str.split('/'))
            else:
                continue  # Skip if date format doesn't match either pattern
                
            shift_date = date(year, month, day)
            
            if employee_name:
                shifts.append({
                    'employee': employee_name,
                    'date': shift_date,
                    'day_of_week': day_of_week,
                    'shift_type': "On-Call Specialty",  # Will be updated when adding to all_shifts
                })
                
        except Exception as e:
            st.error(f"Error parsing row in specialty on-call table {row}: {e}")
            continue
    
    return shifts

def create_calendar_for_employee(shifts, employee_name, cath_lab_shifts=None, ep_shifts=None):
    """Create an iCalendar file with all-day events for a specific employee."""
    # Filter shifts for this specific employee
    employee_shifts = [s for s in shifts if s['employee'].lower() == employee_name.lower()]
    
    # Also check if the employee has any cath lab or EP shifts
    employee_cath_lab_shifts = []
    employee_ep_shifts = []
    
    if cath_lab_shifts:
        employee_cath_lab_shifts = [s for s in cath_lab_shifts if s['employee'].lower() == employee_name.lower()]
        
    if ep_shifts:
        employee_ep_shifts = [s for s in ep_shifts if s['employee'].lower() == employee_name.lower()]
    
    if not employee_shifts and not employee_cath_lab_shifts and not employee_ep_shifts:
        st.warning(f"No shifts found for employee: {employee_name}")
        return None
    
    cal = Calendar()
    cal.add('prodid', '-//Employee Shift Calendar//example.com//')
    cal.add('version', '2.0')
    cal.add('calscale', 'GREGORIAN')
    
    # Group shifts by date to combine multiple shifts on the same day
    shifts_by_date = {}
    
    # Add regular shifts to the grouping
    for shift in employee_shifts:
        date_key = shift['date'].isoformat()
        if date_key not in shifts_by_date:
            shifts_by_date[date_key] = []
        shifts_by_date[date_key].append(shift)
    
    # Add cath lab shifts if they don't overlap with existing dates
    for shift in employee_cath_lab_shifts:
        date_key = shift['date'].isoformat()
        if date_key not in shifts_by_date:
            shifts_by_date[date_key] = []
        shifts_by_date[date_key].append(shift)
    
    # Add EP shifts if they don't overlap with existing dates
    for shift in employee_ep_shifts:
        date_key = shift['date'].isoformat()
        if date_key not in shifts_by_date:
            shifts_by_date[date_key] = []
        shifts_by_date[date_key].append(shift)
    
    # Create events for each date, combining shift information
    for date_key, date_shifts in shifts_by_date.items():
        event = Event()
        
        # Combine all shift types for the summary
        shift_types = [s['shift_type'] for s in date_shifts]
        day_of_week = date_shifts[0]['day_of_week']  # They all have the same date
        shift_date = date_shifts[0]['date']
        
        # Format the summary to show all shift types
        summary = f"{', '.join(shift_types)} - {day_of_week}"
        event.add('summary', summary)
        
        # All-day events need a DATE value type
        event.add('dtstart', shift_date)
        
        # For all-day events, the end date should be the next day
        # The end date is non-inclusive in the iCalendar spec
        end_date = shift_date + timedelta(days=1)
        event.add('dtend', end_date)
        
        event.add('dtstamp', datetime.now())
        
        # Generate a unique ID for the event
        uid = f"{employee_name.replace(' ', '')}-{shift_date.strftime('%Y%m%d')}@shifts.example.com"
        event.add('uid', uid)
        
        # Add description with details about all employees working that day
        description_parts = [f"Your shifts: {', '.join(shift_types)}"]
        
        # Find all employees working on this date
        coworkers_info = []
        for s in shifts:
            # If it's the same date but not the current employee
            if s['date'] == shift_date and s['employee'].lower() != employee_name.lower():
                coworkers_info.append(f"{s['employee']}: {s['shift_type']}")
        
        # Add coworkers section if any exist
        if coworkers_info:
            description_parts.append("\nCoworkers on this day:")
            for info in sorted(coworkers_info):
                description_parts.append(f"- {info}")
        else:
            description_parts.append("\nNo other employees scheduled on this day.")
        
        # Add Cath Lab on-call information if available
        if cath_lab_shifts:
            cath_lab_employee = None
            for shift in cath_lab_shifts:
                if shift['date'] == shift_date and shift['employee'].lower() != employee_name.lower():
                    cath_lab_employee = shift['employee']
                    break
            
            if cath_lab_employee:
                description_parts.append(f"\nCath Lab On-Call: {cath_lab_employee}")
        
        # Add Electrophysiology on-call information if available
        if ep_shifts:
            ep_employee = None
            for shift in ep_shifts:
                if shift['date'] == shift_date and shift['employee'].lower() != employee_name.lower():
                    ep_employee = shift['employee']
                    break
            
            if ep_employee:
                description_parts.append(f"\nElectrophysiology On-Call: {ep_employee}")
        
        event.add('description', "\n".join(description_parts))
        
        cal.add_component(event)
    
    # Return the calendar data
    return cal.to_ical()

def extract_month_year_from_filename(filename):
    """Attempt to extract month and year from the filename."""
    # Example: "ΕΦΗΜΕΡΙΕΣ ΜΑΡΤΙΟΣ 2025.docx"
    month_dict = {
        "ΙΑΝΟΥΑΡΙΟΣ": 1, "ΦΕΒΡΟΥΑΡΙΟΣ": 2, "ΜΑΡΤΙΟΣ": 3, "ΑΠΡΙΛΙΟΣ": 4,
        "ΜΑΙΟΣ": 5, "ΙΟΥΝΙΟΣ": 6, "ΙΟΥΛΙΟΣ": 7, "ΑΥΓΟΥΣΤΟΣ": 8,
        "ΣΕΠΤΕΜΒΡΙΟΣ": 9, "ΟΚΤΩΒΡΙΟΣ": 10, "ΝΟΕΜΒΡΙΟΣ": 11, "ΔΕΚΕΜΒΡΙΟΣ": 12
    }
    
    # Also look for month name in the document content
    month_from_content = None
    if "ΜΑΡΤΙΟΣ" in filename:
        month_from_content = 3
    
    # Default to current month and year if extraction fails
    default_month = datetime.now().month
    default_year = datetime.now().year
    
    try:
        # Try to extract month name and year
        for month_name, month_num in month_dict.items():
            if month_name in filename:
                # Found month, now look for year
                year_match = re.search(r'20\d\d', filename)
                if year_match:
                    year = int(year_match.group())
                    return month_num, year
                return month_num, default_year
        
        # If we found month in content, use that
        if month_from_content:
            year_match = re.search(r'20\d\d', filename)
            if year_match:
                year = int(year_match.group())
                return month_from_content, year
            return month_from_content, default_year
            
    except:
        pass
    
    return default_month, default_year

def download_button(object_to_download, download_filename, button_text):
    """
    Generates a link to download the given object.
    """
    if isinstance(object_to_download, bytes):
        b64 = base64.b64encode(object_to_download).decode()
    else:
        b64 = base64.b64encode(object_to_download.encode()).decode()

    dl_link = f'<a href="data:file/txt;base64,{b64}" download="{download_filename}">{button_text}</a>'
    return dl_link

def main():
    st.title("🏥 Employee Shift Calendar Generator")
    st.subheader("Convert shift schedules to calendar files (.ics)")
    
    st.markdown("""
    This application helps medical staff convert their shift schedules from Word documents into calendar files 
    that can be imported into Google Calendar, Outlook, Apple Calendar, and other calendar applications.
    
    **How to use:**
    1. Upload your main shift schedule document (.docx format)
    2. Optionally upload specialty schedules (Cath Lab, Electrophysiology)
    3. Select the month and year for the schedule
    4. Choose an employee to generate a calendar for
    5. Download the generated calendar file (.ics)
    """)
    
    # File Upload Section
    st.header("1. Upload Schedule Documents")
    
    main_file = st.file_uploader("Upload Main Shift Schedule (.docx)", type=["docx"])
    
    col1, col2 = st.columns(2)
    
    with col1:
        include_cath_lab = st.checkbox("Include Cath Lab on-call shifts?")
        cath_lab_file = None
        if include_cath_lab:
            cath_lab_file = st.file_uploader("Upload Cath Lab On-Call Schedule (.docx)", type=["docx"])
    
    with col2:
        include_ep = st.checkbox("Include Electrophysiology on-call shifts?")
        ep_file = None
        if include_ep:
            ep_file = st.file_uploader("Upload Electrophysiology On-Call Schedule (.docx)", type=["docx"])
    
    # Process files if uploaded
    if main_file is not None:
        st.header("2. Schedule Period")
        
        # Extract month and year from filename if possible
        filename = main_file.name
        detected_month, detected_year = extract_month_year_from_filename(filename)
        
        # Allow user to override detected month/year
        col1, col2 = st.columns(2)
        
        with col1:
            month = st.selectbox(
                "Select Month", 
                options=range(1, 13),
                index=detected_month - 1,  # Adjust for 0-based index
                format_func=lambda x: datetime(2000, x, 1).strftime("%B")
            )
        
        with col2:
            year = st.number_input("Enter Year", min_value=2000, max_value=2100, value=detected_year)
        
        # Process button
        if st.button("Process Schedule Files"):
            with st.spinner("Processing files..."):
                # Read and parse the main document
                tables = read_docx_tables(main_file)
                
                if not tables:
                    st.error("No tables found in the main document.")
                else:
                    # Parse shifts from both tables
                    all_shifts = []
                    
                    # Process first table (if exists)
                    if len(tables) >= 1:
                        first_table_shifts = parse_first_table(tables[0], month, year)
                        all_shifts.extend(first_table_shifts)
                        st.write(f"Found {len(first_table_shifts)} shifts in first table")
                    
                    # Process second table (if exists)
                    if len(tables) >= 2:
                        second_table_shifts = parse_second_table(tables[1], month, year)
                        all_shifts.extend(second_table_shifts)
                        st.write(f"Found {len(second_table_shifts)} shifts in second table")
                    
                    if not all_shifts:
                        st.error("No shifts found in any table!")
                    else:
                        # Process Cath Lab shifts
                        cath_lab_shifts = None
                        if include_cath_lab and cath_lab_file:
                            cath_lab_tables = read_docx_tables(cath_lab_file)
                            if cath_lab_tables:
                                cath_lab_shifts = []
                                for table in cath_lab_tables:
                                    cath_shifts = parse_specialty_on_call_table(table)
                                    for shift in cath_shifts:
                                        shift['shift_type'] = "Cath Lab On-Call"
                                    cath_lab_shifts.extend(cath_shifts)
                                st.write(f"Found {len(cath_lab_shifts)} Cath Lab on-call shifts")
                        
                        # Process Electrophysiology shifts
                        ep_shifts = None
                        if include_ep and ep_file:
                            ep_tables = read_docx_tables(ep_file)
                            if ep_tables:
                                ep_shifts = []
                                for table in ep_tables:
                                    electro_shifts = parse_specialty_on_call_table(table)
                                    for shift in electro_shifts:
                                        shift['shift_type'] = "Electrophysiology On-Call"
                                    ep_shifts.extend(electro_shifts)
                                st.write(f"Found {len(ep_shifts)} Electrophysiology on-call shifts")
                        
                        # Store processed data in session state
                        st.session_state['all_shifts'] = all_shifts
                        st.session_state['cath_lab_shifts'] = cath_lab_shifts
                        st.session_state['ep_shifts'] = ep_shifts
                        st.session_state['processing_complete'] = True
                        
                        # Get unique employee names across all shifts
                        all_employees = sorted(set(shift['employee'] for shift in all_shifts))
                        st.session_state['all_employees'] = all_employees
                        
                        st.success(f"Successfully processed {len(all_shifts)} shift assignments for {len(all_employees)} employees!")
                        st.balloons()
        
        # Generate calendar section - only show if processing is complete
        if st.session_state.get('processing_complete'):
            st.header("3. Generate Calendar")
            
            # Get the list of employees
            all_employees = st.session_state.get('all_employees', [])
            
            # Create a dropdown to select employee
            selected_employee = st.selectbox(
                "Select employee to generate calendar for:",
                options=["All Employees"] + all_employees
            )
            
            if st.button("Generate Calendar"):
                with st.spinner("Generating calendar..."):
                    all_shifts = st.session_state.get('all_shifts', [])
                    cath_lab_shifts = st.session_state.get('cath_lab_shifts')
                    ep_shifts = st.session_state.get('ep_shifts')
                    
                    if selected_employee == "All Employees":
                        # Create a zip file with calendars for all employees
                        st.error("Bulk download for all employees is not implemented yet. Please select a specific employee.")
                    else:
                        # Generate calendar for the selected employee
                        calendar_data = create_calendar_for_employee(
                            all_shifts, 
                            selected_employee, 
                            cath_lab_shifts, 
                            ep_shifts
                        )
                        
                        if calendar_data:
                            # Create file name
                            file_name = f"{selected_employee.replace(' ', '_')}_shifts.ics"
                            
                            # Display download button
                            st.markdown(
                                download_button(
                                    calendar_data, 
                                    file_name, 
                                    f"Download Calendar for {selected_employee}"
                                ),
                                unsafe_allow_html=True
                            )
                            
                            st.success(f"Calendar for {selected_employee} generated successfully!")
                        else:
                            st.error(f"No calendar could be generated for {selected_employee}. No shifts found.")

# Initialize session state variables if they don't exist
if 'processing_complete' not in st.session_state:
    st.session_state['processing_complete'] = False
if 'all_shifts' not in st.session_state:
    st.session_state['all_shifts'] = []
if 'cath_lab_shifts' not in st.session_state:
    st.session_state['cath_lab_shifts'] = None
if 'ep_shifts' not in st.session_state:
    st.session_state['ep_shifts'] = None
if 'all_employees' not in st.session_state:
    st.session_state['all_employees'] = []

# Run the app
if __name__ == "__main__":
    main()
